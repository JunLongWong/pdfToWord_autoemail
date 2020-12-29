# !pip install pdfminer
from pdfminer.layout import LAParams, LTTextBox
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator

import re
from docx import Document
from docx.shared import Cm, Inches
from docx.shared import Pt

from datetime import date, datetime
import datetime
import sys
import pytz

import sensitiveInfomation #sensitive file - not shared on github 
import senderParticulars  #sensitive file - not shared on github 

import email, smtplib, ssl 
from email import encoders 
from email.mime.base import MIMEBase 
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText

def sendemail(fn,receiverEmail):
    sender_email = senderParticulars.email
    receiver_email = receiverEmail
    #email subject
    subject = senderParticulars.subject
    # string to store the body of the mail 
    body = senderParticulars.emailContent
    #sender email password stored in a separate file
    password = senderParticulars.pw
    # instance of MIMEMultipart 
    msg = MIMEMultipart()   
    # storing the senders email address   
    msg['From'] = sender_email 
    # storing the receivers email address  
    msg['To'] = receiver_email  
    # storing the subject  
    msg['Subject'] = subject
    msg["Cc"] = senderParticulars.workEmail
      
    # attach the body with the msg instance 
    msg.attach(MIMEText(body, 'plain')) 
      
    # open the file to be sent  
    filename = fn 

    #Open the file in binary mode
    with open(filename, "rb") as attachment:
        #Add file as application/octet-stream
        #Email client can usually download this automatically as attachment
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())

    #encode file in ASCII characters to send by email
    encoders.encode_base64(part)
    # Add header as key/value pair to attachment part
    part.add_header(
        "Content-Disposition",
        f"attachment; filename= {filename}",
    )

    # Add attachment to message and convert message to string
    msg.attach(part)
    text = msg.as_string()

    # Log in to server using secure context and send email
    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(sender_email, password)
        server.sendmail(sender_email, (receiver_email,senderParticulars.workEmail), text)

#set columns width
def set_col_widths(table):
    widths = (Inches(1), Inches(1.2), Inches(3.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

#Works order Commencement & Delivery date
def cAndD_Date(date):
  while True:
    #test if date is within monday - friday
    if(date.weekday() < 5):
      break
    else:
      date += datetime.timedelta(days=1)
  return date.strftime("%d/%m/%Y")

#Extracting data from PDF by using pre-defined coordinates as reference
def extractData(fileName):
  fp = open(fileName, 'rb')
  rsrcmgr = PDFResourceManager()
  laparams = LAParams()
  device = PDFPageAggregator(rsrcmgr, laparams=laparams)
  interpreter = PDFPageInterpreter(rsrcmgr, device)
  pages = PDFPage.get_pages(fp)

  sorlist=[]        # SOR number
  sordeslist=[]     # SOR Description
  qtyList = []      # Quantity per SOR item
  address_list = [] # Store Blk,Street name, unit no.
  for page in pages:
      interpreter.process_page(page)
      layout = device.get_result()
      for lobj in layout:
                                              # x-coordinates of the pdf document
          if isinstance(lobj, LTTextBox) and (lobj.bbox[0] == 35.6 or lobj.bbox[0] == 357.6 or lobj.bbox[0] == 40.4):
              text = lobj.get_text()
              matchSorNum = re.search(r'\(\d{1} \d+ \)', text)

              matchQty = re.finditer(r'\d+[.]\d+ (m²|FLT|JOB|m|no\.)',text)
              #----getting the address in the WO----
              matchAddress = re.search(r'Reference: Blk.+', text)
              if(matchAddress != None):
                address = matchAddress.group()
                blockIndex  = address.find("Blk")+4
                streetNameIndex = address.find(" ",blockIndex)+1
                unitNoIndex = address.find("#") 
                blockNum = address[blockIndex:streetNameIndex]    # Blk No
                streetName = address[streetNameIndex:unitNoIndex] # Street Name
                unitNo = address[unitNoIndex:]                    # Unit Number
                address_list.append(blockNum)
                address_list.append(streetName)
                address_list.append(unitNo)

              for i in matchQty:
                qtyList.append(i.group())
              #find the newline character, tailor according to data analysed from the PDF file
              newLineChar_pos = text.find("\n")
              if matchSorNum != None:
                # (SOR no.), SOR item description, Quantity
                if (matchSorNum.group() =='(0 5001 )'): # business rule - hardcoded
                  sor_Number = "( " + matchSorNum.group()[3:] # business rule - hardcoded
                else:
                  sor_Number = re.sub(r'0', '1', matchSorNum.group(), 1) # business rules - to replace the 1st SOR number by '1'
                sorlist.append(sor_Number) 
                sordeslist.append(text[matchSorNum.end():newLineChar_pos])

  data=list(zip(sorlist,sordeslist,qtyList))
  return data,address_list

#sor item details
def sorItem(today, document, pdfFile_data):
  # Standard template headers for instruction orders -------------
  items = (
    ("", "Commencement Date", cAndD_Date(today + datetime.timedelta(days=10)), "Delivery Date", cAndD_Date(today + datetime.timedelta(days=25)), ""),
    ("", "Completion Date", cAndD_Date(today + datetime.timedelta(days=25)), "", "", ""),
  )

  # add table 
  table_header = document.add_table(rows=1, cols=6,style="Table Grid")

  #populate header row 
  heading_cells = table_header.rows[0].cells
  heading_cells[0].text = ""
  heading_cells[1].text = 'Flat type'
  heading_cells[2].text = sensitiveInfomation.flat          
  heading_cells[3].text = 'Contract No'
  heading_cells[4].text = sensitiveInfomation.contractNo     
  heading_cells[5].text = ""

  # add a data row for each item
  for i in items:
      cells = table_header.add_row().cells
      cells[0].text = i[0]
      cells[1].text = i[1]
      cells[2].text = i[2]
      cells[3].text = i[3]
      cells[4].text = i[4]
      cells[5].text = i[5]

  # add table ------------------
  items_table = document.add_table(rows=1, cols=6,style="Table Grid")

  # populate header row --------
  heading_cells = items_table.rows[0].cells
  heading_cells[0].text = "S/N"
  heading_cells[1].text = 'Description of works (eg. Location)'
  heading_cells[2].text = 'SOR'
  heading_cells[3].text = 'Quantity'
  heading_cells[4].text = 'Job code'
  heading_cells[5].text = 'Tax code'

  sn = 1
  # populate extracted data from PDF file to the correct column in the table in the word file 
  for sorNumber,sorDescription,qty in pdfFile_data[0]:
      cells = items_table.add_row().cells
      cells[0].text = str(sn)
      cells[1].text = sorDescription
      cells[2].text = sorNumber
      cells[3].text = qty
      cells[4].text = sensitiveInfomation.jobcode              
      cells[5].text = sensitiveInfomation.taxcode                
      sn+=1

def generateDocument(document, pdfFile_data):
  eeName = senderParticulars.name

  p1 = document.add_paragraph()
  p1_word = sensitiveInfomation.WoHeader
  runner_p1 = p1.add_run(p1_word)
  runner_p1.bold = True
  runner_p1.underline =True
  runner_p1.font.name = 'Arial'
  runner_p1.font.size = Pt(16)

  p2 = document.add_paragraph()
  p2_word = "Submitted by Name/Designation: "

  runner_p2 = p2.add_run(p2_word)
  runner_p2.bold = True
  runner_p2_1 = p2.add_run(eeName)
  runner_p2_1.bold = True
  runner_p2_1.underline = True
  fontP2 = runner_p2.font
  fontP2_1 = runner_p2_1.font
  fontP2.size = Pt(12)
  fontP2_1.size = Pt(12)

  p3 = document.add_paragraph()
  p3_word = "Signature: "
  runner_p3 = p3.add_run(p3_word)
  runner_p3.bold = True
  runner_p3_1 = p3.add_run(eeName)
  runner_p3_1.bold = True
  runner_p3_1.underline = True
  font_p3 = runner_p3.font
  font_p3.size = Pt(12)
  font_p3_1 = runner_p3_1.font
  font_p3_1.size = Pt(12)

  p4 = document.add_paragraph()
  sg = pytz.timezone("Asia/Singapore")
  today = datetime.datetime.now(sg)
  today_date = today.strftime("%d/%m/%Y")
  p4_word = "Date: " + str(today_date)
  runner_p4 = p4.add_run(p4_word)
  runner_p4.bold = True
  font_p4 = runner_p4.font
  font_p4.size = Pt(12)

  table = document.add_table(rows=2, cols=3, style="Table Grid")
  set_col_widths(table)

  Blk = table.cell(0,0).paragraphs[0].add_run("Blk")
  Blk.bold=True
  Blk.font.size = Pt(12)
  uN = table.cell(0,1).paragraphs[0].add_run("Unit No.")
  uN.bold=True
  uN.font.size = Pt(12)
  sN = table.cell(0,2).paragraphs[0].add_run("Street Name")
  sN.bold=True
  sN.font.size = Pt(12)
  blk = pdfFile_data[1][0].rstrip()
  unitNo = pdfFile_data[1][2].rstrip()
  sn = pdfFile_data[1][1].rstrip()
  table.cell(1,0).text = blk    #Blk
  table.cell(1,1).text = unitNo #Unit No
  table.cell(1,2).text = sn     #Street Name

  # WO header title
  p5 = document.add_paragraph()
  runner_p5 = p5.add_run(sensitiveInfomation.WoTitle)              
  runner_p5.font.size = Pt(14)

  sorItem(today, document, pdfFile_data)

  filename = blk + "_" + sn + "_" + unitNo + "_ironmongery" +".docx"
  document.save(filename)
  return filename

def main():
  print("Starting to Extract PDF data into word file...\n")
  document=Document()
  pdfData = extractData(sys.argv[1])
  fileName = generateDocument(document, pdfData)
  blk = pdfData[1][0].strip()   # block number
  blk_regex = re.match("(^165a$|^165A$|^9$|^632a$|^632A$|^632b$|^632B$|^536$|^537$)",blk) #business rules

  try:
    if (blk_regex == None):
      raise Exception("Please check & amend the block manually!\n\n")
  except Exception as e:
    print(e)
    print("Email not send!!...Exiting program.")
    sys.exit(0)
  else:
    #Previewing content - to make sure that the text recognition is correct before sending out of email
    print("Previewing content of generated word document\n")
    for sorNo,sorDesc,qty in pdfData[0]:
      print("SOR:{}, Qty:{}, SOR_DESCRIPTION:{}".format(sorNo,qty,sorDesc))

    sendEmail = input("\nConfirm send email(Y/N):")
    if (sendEmail == 'Y' and blk_regex != None):
      print("Proceeding to send email with attached word document....")
      if (blk_regex.group()[:3] == "632"):   #business rule - sending to a different email addresses
        sendemail(fileName,senderParticulars.receiver1Email) 
        print("Email sent to :" + senderParticulars.receiver1Email + " !!!")
      else:
        sendemail(fileName,senderParticulars.receiver2Email) 
        print("Email sent to: " + senderParticulars.receiver2Email + " !!!")
    else:
      print("Word document generated successfully. Email NOT SEND !! Please vet through SOR Items.")

if __name__ == "__main__":
  main()